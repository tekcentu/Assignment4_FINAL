"""Build report/final_report.docx from final_report.md.

The Markdown is hand-written for this report; this script parses just the
constructs it actually contains (ATX headings, paragraphs, bullet lists,
fenced code blocks, GFM tables, figure-image lines, italic captions).
Anything more exotic is out of scope.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


SRC = Path("report/final_report.md")
OUT = Path("report/final_report.docx")
FIG_DIR = Path("report/figures")


# ── Small style helpers ─────────────────────────────────────────────────

def _ensure_style(doc, name, *, font="Calibri", size=11, bold=False,
                  color=None, kind=WD_STYLE_TYPE.PARAGRAPH):
    if name in doc.styles:
        return doc.styles[name]
    st = doc.styles.add_style(name, kind)
    f = st.font
    f.name = font; f.size = Pt(size); f.bold = bold
    if color is not None:
        f.color.rgb = RGBColor.from_string(color)
    return st


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def _add_runs_with_inline(p, text):
    """Honour **bold**, *italic*, and `code` runs inside a paragraph."""
    parts = re.split(
        r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*\n]+\*(?!\*))",
        text,
    )
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif (part.startswith("*") and part.endswith("*")
              and not part.startswith("**")):
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            p.add_run(part)


# ── Markdown parser (subset) ───────────────────────────────────────────

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
FIG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\-\s\|]+\s*$")


def _tokenize(md: str):
    """Yield (kind, payload) tokens."""
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        # fenced code
        if line.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            yield ("code", "\n".join(block))
            continue
        # heading
        m = HEADING_RE.match(line)
        if m:
            yield ("h" + str(len(m.group(1))), m.group(2).strip())
            i += 1
            continue
        # horizontal rule
        if line.strip() == "---":
            yield ("hr", None); i += 1; continue
        # figure
        m = FIG_RE.match(line)
        if m:
            yield ("img", (m.group(1), m.group(2)))
            i += 1
            continue
        # table: header | sep | rows
        if "|" in line and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            yield ("table", (header, rows))
            continue
        # bullet list
        if line.lstrip().startswith(("* ", "- ")):
            items = []
            while i < n:
                lst = lines[i].lstrip()
                if lst.startswith(("* ", "- ")):
                    items.append(lst[2:].rstrip())
                    i += 1
                elif (lines[i].startswith(("  ", "\t"))
                      and lines[i].strip() and items):
                    items[-1] = items[-1] + " " + lines[i].strip()
                    i += 1
                else:
                    break
            yield ("ul", items)
            continue
        # blank
        if not line.strip():
            yield ("blank", None); i += 1; continue
        # paragraph: collect until blank
        para = [line]
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].startswith("#")
            or lines[i].startswith("![")
            or lines[i].startswith("```")
            or lines[i].lstrip().startswith(("* ", "- "))
            or lines[i].strip() == "---"
            or "|" in lines[i] and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1])
        ):
            para.append(lines[i]); i += 1
        yield ("p", " ".join(s.strip() for s in para))


# ── Renderer ───────────────────────────────────────────────────────────

def build():
    doc = Document()
    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Base styles
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(11)
    for lvl, sz, bold in [(1, 20, True), (2, 14, True), (3, 12, True), (4, 11, True)]:
        sty = doc.styles[f"Heading {lvl}"]
        sty.font.name = "Calibri"; sty.font.size = Pt(sz); sty.font.bold = bold
        sty.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    _ensure_style(doc, "Caption-Italic", size=10)
    _ensure_style(doc, "Code-Mono", font="Consolas", size=10)

    md = SRC.read_text(encoding="utf-8")
    tokens = list(_tokenize(md))

    first_h1_done = False
    for kind, payload in tokens:
        if kind == "h1":
            if not first_h1_done:
                # Title page
                first_h1_done = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(payload); r.bold = True; r.font.size = Pt(22)
                r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            else:
                doc.add_heading(payload, level=1)
        elif kind in ("h2", "h3", "h4"):
            doc.add_heading(payload, level=int(kind[1]))
        elif kind == "p":
            if not first_h1_done:
                continue
            txt = payload
            stripped = txt.lstrip()
            is_caption = (
                stripped.startswith("*Figure ")
                or stripped.startswith("*Table ")
                or stripped.startswith("*End of report*")
            )
            if is_caption:
                p = doc.add_paragraph(style="Caption-Italic")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_runs_with_inline(p, txt)
            else:
                p = doc.add_paragraph()
                _add_runs_with_inline(p, txt)
        elif kind == "ul":
            for it in payload:
                p = doc.add_paragraph(style="List Bullet")
                _add_runs_with_inline(p, it)
        elif kind == "code":
            p = doc.add_paragraph(style="Code-Mono")
            p.add_run(payload)
            # light grey shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4"); pPr.append(shd)
        elif kind == "img":
            alt, path = payload
            img = Path(path)
            if not img.is_absolute():
                img = Path("report") / img if not img.exists() else img
            if not img.exists():
                # try relative to figures dir
                img = FIG_DIR / Path(path).name
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img), width=Inches(6.2))
            else:
                doc.add_paragraph(f"[missing image: {path}]")
        elif kind == "table":
            header, rows = payload
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                c = tbl.rows[0].cells[j]
                c.text = ""
                _shade(c, "1F3A5F")
                p = c.paragraphs[0]; r = p.add_run(h); r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for i, row in enumerate(rows, start=1):
                for j, cell in enumerate(row):
                    c = tbl.rows[i].cells[j]
                    c.text = ""
                    p = c.paragraphs[0]
                    _add_runs_with_inline(p, cell)
                    for run in p.runs:
                        run.font.size = Pt(10)
        elif kind == "hr":
            doc.add_paragraph()
        elif kind == "blank":
            continue

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
